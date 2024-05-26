import os

from docx import Document
from dotenv import load_dotenv

from celi_framework.utils.llms import quick_ask
from celi_framework.utils.utils import load_text_file, load_json
from templates.templates import make_gpt_prompt_simple, score_llm_prompt

load_dotenv()

gre_dir = os.path.dirname(__file__)
CURRENT_QUESTIONS_DIR = f"{gre_dir}/data/working_data/targets"
EXAMPLE_QUESTIONS_DIR = f"{gre_dir}/data/working_data/examples"
EXAMPLE_RUBRIC_DIR = f"{gre_dir}/data/working_data/rubric"

schema = load_json(f"{gre_dir}/../AP_English_Language/schema.json")

# for i in range(1, 2):
#     for key in schema.keys():


def run_gpt():
    for key in ["q1", "q2", "q3"]:

        current_prompt = load_text_file(
            file_path=f"{CURRENT_QUESTIONS_DIR}/set_3_target_{key}_question.txt"
        )

        example_prompt = load_text_file(
            file_path=f"{EXAMPLE_QUESTIONS_DIR}/set_1_example_q1_question.txt"
        )
        example_response = load_text_file(
            file_path=f"{EXAMPLE_QUESTIONS_DIR}/set_1_example_q1_response.txt"
        )
        example_rubric = load_text_file(
            file_path=f"{EXAMPLE_RUBRIC_DIR}/gre_1_scoring_guidelines.txt"
        )

        example_prompt_response = f"Prompt:\n{example_prompt}\n\n{example_response}"

        system_message, user_prompt = make_gpt_prompt_simple(
            current_prompt=current_prompt,
            example_prompt_response=example_prompt_response,
        )

        whole_prompt = f"{system_message}\n\n{user_prompt}"

        print("Whole Prompt:")
        print(f"{whole_prompt}")

        # print(f"{whole_prompt}\n\n")
        llm_response = quick_ask(
            prompt=whole_prompt, token_counter=None, model_name="gpt-4o"
        )

        # save gpt_outputs
        print("successfully created gpt output for question " + key)
        print(llm_response)
        # Specify the directory and filename
        file_path = f"output/gpt/{key}_final_essay_gpt_unscored.docx"

        # Create a new Document
        doc = Document()
        doc.add_paragraph(llm_response)  # Adding the text as a new paragraph

        # Save the document to the specified file path
        doc.save(file_path)

        system_message, user_prompt = score_llm_prompt(
            prompt=current_prompt, llm_output=llm_response, rubric=example_rubric
        )

        whole_prompt = f"{system_message}\n\n{user_prompt}"

        print("Whole Prompt:")
        print(f"{whole_prompt}")

        # print(f"{whole_prompt}\n\n")
        llm_response = quick_ask(prompt=whole_prompt, token_counter=None)

        # Specify the directory and filename
        file_path = f"output/gpt/{key}_final_essay_gpt_scored.docx"

        # Create a new Document
        doc = Document()
        doc.add_paragraph(llm_response)  # Adding the text as a new paragraph

        # Save the document to the specified file path
        doc.save(file_path)


run_gpt()
