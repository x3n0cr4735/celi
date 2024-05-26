# TODO -> Put stuff in utils/common where it makes sense

from dotenv import load_dotenv

load_dotenv()

import json
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Dict

from celi_framework.core.job_description import ToolImplementations
from celi_framework.utils.log import app_logger as logger
from celi_framework.utils.utils import load_json, load_text_file
from docx import Document

cur_dir = os.path.dirname(os.path.realpath(__file__))


@dataclass
class GREToolImplementations(ToolImplementations):
    def __post_init__(self):
        self.schema = load_json(f"{cur_dir}/schema.json")

    # Retrieves the question numbers and topics of the questions that need to be answered.
    def get_schema(self) -> Dict[str, str]:
        return self.schema

    def get_rubric(self) -> str:
        """
        Retrieves the rubric to score the questions.

        Returns:
            str: The rubric to score the question.
        """
        rubric = load_text_file(
            f"{cur_dir}/data/working_data/rubric/gre_1_scoring_guidelines.txt"
        )

        return rubric

    def retrieve_question_prompt(self, question_number: str) -> str:
        """
        Retrieves the prompt for a given question in the free response section of the test.

        Args:
            question_number (str): The unique identifier for the question (q1, q2, or q3).

        Returns:
            str: The prompt for the question.
        """
        prompt = load_text_file(
            f"{cur_dir}/data/working_data/targets/set_3_target_{question_number}_question.txt"
        )

        return prompt

    def retrieve_example_prompt_response(self, question_number: str) -> str:
        """
        Retrieves a prompt, answer pair for a question that tests similar skills as the question being answered in the free response section of the test.

        Args:
            question_number (str): The unique identifier for the question (q1).

        Returns:
            str: Prompt, question example pair.
        """

        prompt = load_text_file(
            f"{cur_dir}/data/working_data/examples/set_1_example_q1_question.txt"
        )
        response = load_text_file(
            f"{cur_dir}/data/working_data/examples/set_1_example_q1_response.txt"
        )
        pair = f"{prompt}\n\n{response}"

        return pair

    # def retrieve_example_response(self, question_number: str) -> str:
    #     """
    #     Retrieves a prompt, answer pair for a question that tests similar skills as the question being answered in the free response section of the test.
    #
    #     Args:
    #         question_number (str): The unique identifier for the question (q1, q2, or q3).
    #
    #     Returns:
    #         str: Prompt, question example pair.
    #     """
    #     response = load_text_file(f"{ROOT_DIR}/celi_framework/examples/GRE/data/working_data/examples/set_1_example_{question_number}_response.txt")
    #
    #     return response

    def retrieve_instructions(self, question_number: str) -> str:
        """
        Retrieves a prompt, answer pair for a question that tests similar skills as the question being answered in the free response section of the test.

        Args:
            question_number (str): The unique identifier for the questions (q1,q2,or q3).

        Returns:
            str: Prompt, question example pair.
        """
        instructions = load_text_file(
            f"{cur_dir}/data/working_data/instructions/set_1_target_q1_instructions.txt"
        )

        return instructions

    def save_draft(self, draft_dict: str) -> dict:
        """
        Saves the provided draft as a JSON file and returns a structured dictionary.
        Additionally, signals via Redis that a draft is ready for download.

        Args:
            draft_dict (str): A string representation of a dictionary containing the draft text under the 'draft_dict' key.

        Returns:
            dict: The structured dictionary that was saved.
        """

        question_number = "q1"

        def save_draft_to_docx(draft_dict, question_number, timestamp):
            """
            Saves the draft from a dictionary into a .docx file in the specified directory

            Args:
                draft_dict (dict): The dictionary containing the draft text.
                output_dir_path (str): The directory path where the new .docx file should be saved.
                output_filename (str): The name of the file to be saved.
            """
            # Extract the draft text
            # draft_text = draft_dict.get('draft', '')

            output_docx_path = os.path.join(
                cur_dir, "output/word", f"{question_number}-response-{timestamp}.docx"
            )

            doc = Document()
            # Create a new .docx document and save it
            # Iterate through each key-value pair in the dictionary
            for key, value in draft_dict.items():
                # Add a heading or bold paragraph for the section title
                doc.add_heading(
                    key, level=1
                )  # You can adjust the level according to your needs

                # Add the paragraph text from the dictionary
                doc.add_paragraph(value)

            # Save the document to the specified path
            doc.save(output_docx_path)

            return output_docx_path

        timestamp = datetime.now().strftime("%m%d%y-%H%M%S")
        logger.info(
            f"DRAFT DICT TEXT LOOKS LIKE THIS\n{draft_dict}", extra={"color": "cyan"}
        )

        json_path = os.path.join(
            cur_dir, "output/draft_dict", f"{question_number}-response-{timestamp}.json"
        )

        # # Assuming draft_dict is a string that needs to be loaded into a dictionary
        # structured_dict = json.loads(draft_dict)
        # structured_dict = structured_dict.get('draft_dict', '')
        # structured_dict = json.loads(structured_dict)

        try:
            # Assuming draft_dict is a JSON string
            structured_dict = json.loads(draft_dict)

        except json.JSONDecodeError as e:
            print(f"Failed to decode JSON: {e}")
            # Handle the error or fix the data as needed
            return None

        # Define the full path for the saved file
        os.makedirs(
            os.path.dirname(json_path), exist_ok=True
        )  # Ensure directory exists

        # Save the JSON file
        with open(json_path, "w") as json_file:
            json.dump(structured_dict, json_file)

        try:
            # Assuming save_draft_to_docx returns the full path to the saved .docx file
            docx_filepath = save_draft_to_docx(
                structured_dict, question_number, timestamp
            )
            if docx_filepath:
                redis_key = f"draft_ready:{question_number}"
                # redis_client.set(redis_key, docx_filepath)
                logger.info(
                    f"Signaled Redis that draft for question {question_number} is ready.",
                    extra={"color": "cyan"},
                )
            else:
                logger.error(
                    f"Failed to save the .docx for question {question_number}."
                )
        except Exception as e:
            logger.error(f"Error saving draft or signaling Redis: {e}")

        # Update Redis with the status that we're waiting for user feedback
        # redis_client.set(self.progress_key, "Saving the draft")

        # FastAPI: When the draft is ready
        # redis_client.set(f'draft_ready:{question}', 'true')

        return structured_dict
